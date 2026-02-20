"""根据 Git 提交日志生成 Word 更新文档"""
import subprocess
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from collections import OrderedDict
import re

# 获取 git log
result = subprocess.run(
    ["git", "log", "--pretty=format:%h|%ai|%s", "--no-merges"],
    capture_output=True, text=True, encoding="utf-8"
)

# 按日期分组
date_groups = OrderedDict()
for line in result.stdout.strip().split("\n"):
    parts = line.split("|", 2)
    if len(parts) != 3:
        continue
    hash_id, datetime_str, message = parts
    date = datetime_str.split(" ")[0]
    if date not in date_groups:
        date_groups[date] = []
    type_match = re.match(r"^(\w+):\s*(.+)", message)
    if type_match:
        commit_type = type_match.group(1)
        desc = type_match.group(2)
    else:
        commit_type = "other"
        desc = message
    date_groups[date].append({
        "hash": hash_id,
        "type": commit_type,
        "desc": desc,
    })

TYPE_MAP = {
    "feat": "新功能", "fix": "修复", "refactor": "重构", "perf": "性能优化",
    "style": "样式", "chore": "维护", "debug": "调试", "config": "配置",
    "optimize": "优化", "update": "更新", "other": "其他",
}

TYPE_COLORS = {
    "新功能": RGBColor(0x22, 0x8B, 0x22), "修复": RGBColor(0xCC, 0x33, 0x33),
    "重构": RGBColor(0x33, 0x66, 0xCC), "性能优化": RGBColor(0xFF, 0x8C, 0x00),
    "样式": RGBColor(0x99, 0x33, 0xCC), "维护": RGBColor(0x66, 0x66, 0x66),
    "调试": RGBColor(0x99, 0x99, 0x00), "配置": RGBColor(0x00, 0x80, 0x80),
    "优化": RGBColor(0xFF, 0x8C, 0x00), "更新": RGBColor(0x33, 0x66, 0xCC),
    "其他": RGBColor(0x66, 0x66, 0x66),
}

doc = Document()
style = doc.styles["Normal"]
style.font.name = "微软雅黑"
style.font.size = Pt(10.5)

title = doc.add_heading("AI 文生视频平台 — 更新日志", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
total = sum(len(v) for v in date_groups.values())
run = subtitle.add_run(f"记录周期：{list(date_groups.keys())[-1]} ~ {list(date_groups.keys())[0]}  |  共 {total} 次提交")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

for date, commits in date_groups.items():
    doc.add_heading(date, level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        row.cells[0].width = Cm(2.5)
        row.cells[1].width = Cm(2.5)
        row.cells[2].width = Cm(12)
    for i, text in enumerate(["提交ID", "类型", "描述"]):
        r = table.rows[0].cells[i].paragraphs[0].add_run(text)
        r.bold = True
        r.font.size = Pt(10)
    for c in commits:
        row = table.add_row().cells
        r0 = row[0].paragraphs[0].add_run(c["hash"])
        r0.font.size = Pt(9)
        r0.font.name = "Consolas"
        type_cn = TYPE_MAP.get(c["type"].lower(), c["type"])
        r1 = row[1].paragraphs[0].add_run(type_cn)
        r1.font.size = Pt(9)
        r1.bold = True
        r1.font.color.rgb = TYPE_COLORS.get(type_cn, RGBColor(0x33, 0x33, 0x33))
        r2 = row[2].paragraphs[0].add_run(c["desc"])
        r2.font.size = Pt(9)
    doc.add_paragraph("")

# 统计摘要
doc.add_heading("提交统计", level=2)
type_count = {}
for commits in date_groups.values():
    for c in commits:
        t = TYPE_MAP.get(c["type"].lower(), c["type"])
        type_count[t] = type_count.get(t, 0) + 1

st = doc.add_table(rows=1, cols=2)
st.style = "Light Grid Accent 1"
st.alignment = WD_TABLE_ALIGNMENT.CENTER
st.rows[0].cells[0].paragraphs[0].add_run("类型").bold = True
st.rows[0].cells[1].paragraphs[0].add_run("次数").bold = True
for t, count in sorted(type_count.items(), key=lambda x: -x[1]):
    row = st.add_row().cells
    r = row[0].paragraphs[0].add_run(t)
    r.font.color.rgb = TYPE_COLORS.get(t, RGBColor(0x33, 0x33, 0x33))
    r.bold = True
    row[1].paragraphs[0].add_run(str(count))

doc.save("更新日志.docx")
print("文档已生成: 更新日志.docx")
